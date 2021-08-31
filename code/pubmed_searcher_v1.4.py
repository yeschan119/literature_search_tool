#!/usr/bin/env python
# coding: utf-8

# In[11]:


# Import docx NOT python-docx
# date: 2021/07/27
# developer: eungchan-kang
# email: yeschanf119@gmail.com
# company : 대웅제약 pvteam

import time
import pandas as pd
import numpy as np
import requests, bs4
from bs4 import BeautifulSoup
import docx
import platform
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
import nltk
from nltk.tokenize import word_tokenize
from collections import Counter
import matplotlib.pyplot as plt
from matplotlib import rcParams
import os
from os import path
import re
import json
import sys
from io import BytesIO
import math
import tkinter as tk
from tkinter.ttk import*
from tkinter import Tk, ttk, Frame, LabelFrame, filedialog
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QVBoxLayout, QCalendarWidget
from PyQt5.QtCore import QDate
import webbrowser
import threading

from IPython.core.display import display, HTML
display(HTML("<style>.container { width:100% !important; }</style>"))

# pharmacovigilance keywords
pv_keywords = ['patient','serious','year','old','male','female','year','year-old','drug-drug'
            'woman','man','adverse','reaction','adult','side','effect','drug','event']
# keywords in essay
thesis_keyword = ['aims','methods','results','conclusion','keywords'
                 ,'purpose','materials and methods','background','result']

p_name = ['olmesartan OR amlodipine OR rosuvastatin','olmesartan OR rosuvastatin','human AND epithelial AND growth AND factor','recombinant AND human AND erythropoietin',
         'leuprolide AND acetate','risendronate OR cholecalciferol','somatropin','dioctahedral smectite','pitavastatin','entecavir','ursodeoxycholic acid',
         'Biodiastase OR diastase OR protease OR cellulose OR crease-PEG OR dizet-100 OR lipase OR bromelain OR ursodeoxycholic acid OR simethicone OR penprosin OR pancellase OR pancreatin',
         'atorvastatin OR metformin','camostat mesilate, remdesivir','flurbiprofen','choline alfoscerate','aluminum hydroxide OR magnesium hydroxide','tobramycin',
         'erdosteine','cefotiam hydrochloride','meropenem','riboflavin OR folic acid OR thiamine nitrate OR pyridoxine OR calcium pantothenate OR niacinamide OR cyanocobalamin OR zinc oxide',
         'alitretinoin','tenofovir AND alfenamide','clopidogrel','olmesartan OR amlodipine OR rosuvastatin OR ezetimibe','olmesartan OR amlodipine OR rosuvastatin OR ezetimibe',
         'niclosamide','Prolyl-tRNA synthease inhibitor OR pirfenidone OR nintedanib OR nebivolol OR paroxetine','Clostridium botulinum A toxin','alitretinoin','aceclofenac',
         'trafermin','olopatadine AND hydrochloride','mosapride','olmesartan','metformin','fentanyl','fentanyl']

#get path from exe file
def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

root = Tk()
trained_words = {}

# open file that has trained_words from ICSR file
def get_trained_source():
    global trained_words
    try:
        with open("trained_words.json") as json_file:
                trained_words = json.load(json_file)
        return True
    except:
        messagebox.askokcancel("오류!", "소스파일이 존재하지 않습니다.")
        return False

# Natural Language Process

class NLP:
    def __init__(self,URL):
        self.URL = URL
        self.title = ''
        self.contents = ''
        self.NN_words = []
    def web_crawling(self):
        resp = requests.get(self.URL)
        try:
            resp.raise_for_status()  # 200 OK 코드가 아닌 경우 에러 발동
        except:
             messagebox.askokcancel("오류!", "소스파일이 존재하지 않습니다.")
        # resp.status_code : 응답값 받기
        if(resp.status_code == requests.codes.ok):
            html = resp.text
        else:
            print("response error")
        # html parsing 하기
        bs = bs4.BeautifulSoup(html, 'html.parser')
        tags = str(bs.select('title'))  # get title
        self.title = tags.split(',')  # there are several titles
        for link in bs.find_all('div'):
            if link.get('id') == 'abstract':  # class named abstract has main contents
                self.contents += link.text.strip()
        word_tokens = self.contents.replace('-',' ').split(' ')
        i = self.title[0].find('>') + 1  # remove first tag <title>
        j = self.title[0].find('</title>')  # remove second tag </title>
        if j != -1:  # no second tag
            return self.title[0][i:j], word_tokens
        else:
            return self.title[0][i:], word_tokens
        
    def process_data(self):
        cleaned_contents = re.sub(r'[^\.\?\!\w\d\s]','',self.contents)  # remove special characters
        cleaned_contents = cleaned_contents.lower()  # make all words to lower
        word_tokens = word_tokenize(cleaned_contents) # split sentences by token
        token_pos = nltk.pos_tag(word_tokens)  # tag the part of speech to each token
        self.extract_nouns(token_pos)  # only get nouns from tokens
        
    def extract_nouns(self,token_pos):
        for word, pos in token_pos:  # pos : part of speech
            if 'NN' in pos or 'NNS' in pos:  # NN : noun, NNS : nouns
                self.NN_words.append(word)
                
    def show_barchart(self):
        memfile = BytesIO() # build a buffer to store barchart
        counted_words = Counter(self.NN_words)
        words = []
        counts = []
        for letter, count in counted_words.most_common(10):
            words.append(letter)
            counts.append(count)
        # Figure Size
        fig, ax = plt.subplots(figsize =(15, 7))
        # Horizontal Bar Plot
        ax.barh(words, counts)
        # Add padding between axes and labels
        ax.xaxis.set_tick_params(pad = 5)
        ax.yaxis.set_tick_params(pad = 10)

        # Add x, y gridlines
        ax.grid(b = True, color ='grey',
                linestyle ='-.', linewidth = 0.5,
                alpha = 0.2)

        # Show top values
        ax.invert_yaxis()

        # Add annotation to bars
        for i in ax.patches:
            plt.text(i.get_width()+0.1, i.get_y()+0.5,
                     str(round((i.get_width()), 2)),
                     fontsize = 10,
                     color ='grey')

        # Add Plot Title
        ax.set_title('Top10 words in the paper & their count',fontweight='bold',
                     loc ='center',fontsize='14' )
        
        # Add Text watermark
        fig.text(0.9, 0.15, 'PVteam', fontsize = 12,
                 color ='grey', ha ='right', va ='bottom',
                 alpha = 0.7)
        plt.xlabel('number of word', fontweight='bold', fontsize = 12)
        plt.ylabel('word',fontweight='bold',fontsize=12)
        plt.savefig(memfile)
        #path = resource_path('barchart.png')
        # Show Plot
        #plt.savefig(path, bbox_inches='tight')
        return memfile
        #add_new_to_trained_words
        
    def start_learning(self):
        tw_count = np.array(list(trained_words.values()))
        max_count = math.ceil((np.max(tw_count)) / 50)
        for word in self.NN_words:
            if word not in trained_words:
                trained_words[word] = 1
            else:
                trained_words[word] += 1
        temp = {}
        for key in trained_words.keys():
            #해당 단어가 논문 키워드면 안되고, 단어의 길이가 2자 이상이어야 하고, 등장횟수가 max_count / 50 이상이면..
            if key not in thesis_keyword and len(key) > 1 and trained_words[key] >= max_count:
                temp[key] = trained_words[key]
        new_trained_words = dict(sorted(temp.items(), key=lambda item: item[1], reverse = True))
        path = resource_path('trained_words.json')
        with open(path,"w") as outfile:
            json.dump(new_trained_words, outfile)
            
    def output_result(self):
        #self.start_learning()
        return(self.show_barchart())
        
class WordMarker:
    def __init__(self,title, contents):
        self.title = title
        self.contents = contents
        self.kw_sim = 0  # keyword similarity
        self.tw_sim = 0  # trained words similarity
        self.basic_doc = docx.Document()
        self.pro_doc = docx.Document()
        self.basic_stream = BytesIO()  # save doc to buffer
        self.pro_stream = BytesIO()  # save doc to buffer
        
        
    def cleanText(self,readData):
        #텍스트에 포함되어 있는 특수 문자 제거
        text = re.sub('[;=+,#/\?:^$.@*\"※~&%ㆍ!』\\‘|\(\)\[\]\<\>`\'…》]', '', readData)
        text = re.sub(r'[0-9]+', '', text)
        wlem = nltk.WordNetLemmatizer()
        new_word = wlem.lemmatize(text)
        return new_word

    def get_similarity(self):
        similarity = round(self.tw_sim / len(trained_words) * 100)
        if similarity + self.kw_sim < 100:
            return similarity + self.kw_sim
        else:
            return 100

    def write_total_doc(self,memfile):
        tw_count = np.array(list(trained_words.values()))
        mean_count = math.ceil(np.mean(tw_count))
        
        # Add a Title to the document
        self.pro_doc.add_heading(self.title, 0)
        # Creating paragraph with some content
        para = self.pro_doc.add_paragraph()
        for word in self.contents:
            cleanedWord = self.cleanText(word)
            if cleanedWord in pv_keywords:
                para.add_run(word).font.highlight_color = WD_COLOR_INDEX.YELLOW
                if self.kw_sim < 90: # keyword similarity must have limit
                    self.kw_sim += 10 # keyword similarity  
            elif cleanedWord in trained_words and trained_words[cleanedWord] > mean_count:
                para.add_run(word).font.highlight_color = WD_COLOR_INDEX.GRAY_25
                self.tw_sim += 1
            else:
                para.add_run(word)
            para.add_run(' ')
        similarity = self.get_similarity()
        # Now save the document to a location
        self.pro_doc.add_heading('Bar Chart',0)
        self.pro_doc.add_picture(memfile, width=Inches(5.0), height=Inches(3.0))
        #path_pro_doc = resource_path('pro.docx')
        self.pro_doc.save(self.pro_stream)
        memfile.close()
        return similarity
    
    def write_basic_doc(self):
        # Add a Title to the document
        self.basic_doc.add_heading(self.title, 0)
        # Creating paragraph with some content
        para = self.basic_doc.add_paragraph()
        #highlight to the keywords in doc
        for word in self.contents:
            cleanedWord = self.cleanText(word)
            if cleanedWord in pv_keywords:
                para.add_run(word).font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                para.add_run(word)
            para.add_run(' ')
        # Now save the document to a location
        #path_basic_doc = resource_path('basic.docx')
        self.basic_doc.save(self.basic_stream)
        
    def show_doc(self, file, option):
        f = open(file, 'wb')
        if option == 1:
            f.write(self.basic_stream.getvalue())
        else:
            f.write(self.pro_stream.getvalue())
        # open the document with any OS
        if platform.system() == 'Darwin':   # mac OS
            subprocess.call(['open',file])
        elif platform.system() == 'Windows':  # windows
            os.startfile(file)
        else:  # Linux
            subprocess.call(['xdg-open',file])
    def save_doc(self, file, option):
        if option == 1:
            self.basic_doc.save(file)
        else:
            self.pro_doc.save(file)

    
class CalendarApp(QWidget):
    def __init__(self):
        super().__init__()
        self.Date = ''

    def Calendar(self):
        cal = QCalendarWidget(self)
        cal.setGridVisible(True)
        cal.clicked[QDate].connect(self.showDate)

        self.lbl = QLabel(self)
        date = cal.selectedDate()
        self.lbl.setText(date.toString())

        vbox = QVBoxLayout()
        vbox.addWidget(cal)
        vbox.addWidget(self.lbl)

        self.setLayout(vbox)

        self.setWindowTitle('check date')
        #self.setGeometry(200, 200, 300, 200)
        self.show()
    
    def showDate(self, date):
        self.lbl.setText(date.toString())
        #print(date.toString())
        self.Date =  date.toString('yyyy/MM/dd')
    def get_date(self):
        return self.Date
            
# Literature Search GUI class

class LS_GUI:
    def __init__(self,root):
        super().__init__()
        self.URL = tk.StringVar()
        self.choice = 0
        self.similarity = 0
        self.p_var = tk.DoubleVar()
        self.window = root
        self.radioVar = tk.IntVar()  # radiobox to choose only one of them(line listing, summary)
        self.marker = object()
        self.file_path = ''
        self.p_name = tk.StringVar()
        self.p_var = tk.DoubleVar()
        self.s_var = tk.DoubleVar()
        self.From = ''
        self.To = ''
    
    def from_calendar(self):
        app = QApplication(sys.argv)
        start_date = CalendarApp()
        start_date.Calendar()
        app.exec_()
        self.From = start_date.get_date()
        tk.Label(self.window, text = self.From, font = ('sans 11 bold'), width = 10).place(x = 140, y = 147)
        
    def to_calendar(self):
        app = QApplication(sys.argv)
        end_date = CalendarApp()
        end_date.Calendar()
        app.exec_()
        self.To = end_date.get_date()
        tk.Label(self.window, text = self.To, font = ('sans 11 bold'), width = 10).place(x = 140, y = 197)
        
    def open_browser(self):
        search_keyword = '''(((''' + self.p_name.get() + ''') AND ((((("adverse effects" 
            OR "adverse drug reaction" OR "side effect" OR "adverse effect"[Title/Abstract])) 
            OR (("treatment failure" OR "lack of efficacy" OR "nonresponse" OR "no response" 
            OR "unresponse"[Title/Abstract]))) OR (("safety" OR "drug safety" OR "drug toxicity" 
            OR "toxicity"[Title/Abstract]))) OR (("medication error" OR "drug misuse" 
            OR "medication error" OR "drug abuse" OR "drug overdose" OR "death" 
            OR "off label drug use" OR "pregnancy" OR "breast feeding" OR "lactation"))) 
            AND (''' + self.From + '''[date - publication] : ''' + self.To + '''[date - publication])'''
        baseUrl = 'https://pubmed.ncbi.nlm.nih.gov/?term='
        url = baseUrl + search_keyword
        webbrowser.open(url)
        
    # 경고창에서 '확인'을 누르면 종료, '취소'를 누르면 재실행
    def close_program(self):
        if messagebox.askokcancel("오류!", "basic 혹은 pro 하나를 선택하셔야 합니다."):
            self.disable_event()
        else:
            self.window.destroy()
            exit()
            
    def getDir(self, initialDir=""):
        if self.file_path == '':
            if(initialDir != ""):
                self.file_path = filedialog.askdirectory(initialdir=initialDir)
            else:
                self.file_path = filedialog.askdirectory()
        return self.file_path + '/'

    def open_doc(self):
        global marker
        if self.choice == 1:
            open_file = self.getDir() + 'basic.docx'
            #open_file = filedialog.asksaveasfile(mode='w', defaultextension="docx", initialfile = 'basic')
            doc_stream = self.marker.show_doc(open_file,self.choice)
        elif self.choice == 2:
            open_file = self.getDir() + 'pro.docx'
            #open_file = filedialog.asksaveasfile(mode='w', defaultextension="docx", initialfile = 'pro')
            doc_stream = self.marker.show_doc(open_file,self.choice)
        else:
            self.close_program()

    def save_as(self):
        global marker
        fname = asksaveasfilename(filetypes=(("word files", "*.docx",),
                                        ("All files", "*.*")))
        fname += '.docx'
        # note: this will fail unless user ends the fname with ".xlsx"
        if self.choice == 0:
            self.close_program()
        else:
            marker.save_doc(fname, self.choice)
        
    def close_window(self):
        try:
            self.marker.basic_stream.close()
            self.marker.pro_stream.close()
            self.window.destroy()
        except:
            self.window.destroy()
            
    def show_progress(self):
        progress_bar = tk.ttk.Progressbar(self.window, maximum = 100, length=230, variable = self.p_var)
        progress_bar.place(x = 70, y = 485)
        txt = tk.Label(self.window,text = '')
        txt.place(x= 310 ,y=484)
        for i in range(0,101):
            time.sleep(0.01)
            self.p_var.set(i)
            progress_bar.update()
            progress_bar['value'] += 20
            txt['text']=str(int(progress_bar['value']))+'%','완료' 
            
    def show_similarity(self):
        progress = tk.ttk.Progressbar(self.window, maximum = 100, length=230, variable=self.s_var)
        progress.place(x = 70, y = 575)
        txt = tk.Label(self.window,text = '')
        txt.place(x= 310 ,y=574)
        for i in range(0,self.similarity+1):
            time.sleep(0.01)
            self.s_var.set(i)
            progress.update()
            progress['value'] += 20
            txt['text']=str(int(progress['value']))+'%','확신'
    
    def search(self):
        nlp = NLP(self.URL.get())
        title, contents = nlp.web_crawling()
        nlp.process_data()
        memfile = nlp.output_result()
        self.marker = WordMarker(title,contents)
        self.marker.write_basic_doc()
        self.similarity = self.marker.write_total_doc(memfile)
        self.show_progress()
        self.show_similarity()
    # radiovar를 통해 line listing과 summary tabulation을 정하는 함수
    def radCall(self):
        self.choice = self.radioVar.get() # radio 버튼에서 사용자가 클릭한 부분에 해당하는 값이 출력
    
    def disable_event(self):
        pass
    
    def reset_window(self):
        global progress_bar
        progress_bar = 0
        self.close_window()
        root = Tk()
        search = LS_GUI(root)
        search.main_window()
            
    def main_window(self):
        
        self.window.title("pvteam")
        #self.window.overrideredirect(True)
        self.window.geometry('450x720+400+300')
        self.window.resizable(False, False)
        
        #========================================================== 논문 리스트 검색 ==========================================================
        P_Frame = LabelFrame(self.window, text = '논문 리스트 추출',width = 410, height = 290)
        P_Frame['borderwidth'] = 3
        P_Frame['relief'] = 'groove'
        P_Frame.place(x = 20, y = 20)
        tk.Label(self.window, text = '품목명:').place(x = 55, y = 50)
        tk.Entry(self.window, textvariable = self.p_name).place(x=115, y = 50, width=285)
        
        location_combo = Combobox(self.window, width = 12, textvariable = self.p_name)
        location_combo['values'] = p_name
        #location_combo.current(0)
        location_combo.place(x=60, y = 85, width = 340)
        
        
        D_Frame = LabelFrame(self.window, text = '기간 선택',width = 340, height = 120)
        D_Frame['borderwidth'] = 2
        D_Frame['relief'] = 'ridge'
        D_Frame.place(x = 60, y = 120)
        
        tk.Button(self.window, text='from', width = 5, fg = 'black', command = self.from_calendar).place(x = 70, y=147)
        tk.Button(self.window, text='to', width = 5, fg = 'black', command = self.to_calendar).place(x = 70, y = 197)
        
        
        tk.Button(self.window, text = '검색', bg = '#448aff', fg = 'black',
                  font = ('sans 11 bold'), width = 10, command = self.open_browser).place(x = 170, y = 260)
        
        
        #================================================= ICSR/Safety check =======================================================================
        P_Frame = LabelFrame(self.window,text = 'ICSR/Safety 검사',width = 410, height = 365)
        P_Frame['borderwidth'] = 3
        P_Frame['relief'] = 'groove'
        P_Frame.place(x = 20, y = 330)
        
        #self.window.configure(bg='#e8eaf6')
        # 사용자가 해당 버튼을 클릭하면 그 버튼에 해당하는 ragiovar가 1로 설정
        rad1 = tk.Radiobutton(self.window, text = 'basic mode',variable=self.radioVar, value=1, command=self.radCall)
        # columnspan 속성은 해당 위젯이 여러열에 걸쳐져 있다는 뜻
        rad1.place(x = 85, y = 360)
        # 사용자가 해당 버튼을 클릭하면 그 버튼에 해당하는 ragiovar가 2로 설정
        
        rad2 = tk.Radiobutton(self.window, text= 'pro mode',variable=self.radioVar, value=2, command=self.radCall)
        rad2.place(x = 195, y = 360)
        
        tk.Button(self.window, text = "reset", bg = '#e0e0e0', fg = 'black',
                   font = ('sans 8 bold'), width = 6, command = self.reset_window).place(x = 300, y= 360)
        tk.Label(self.window, text = "URL: ").place(x = 40, y = 402)
        tk.Entry(self.window, textvariable = self.URL).place(x = 82, y = 402, width = 310)
        #tk.Label(self.window,text = "시작하시려면 시작 버튼을 클릭하세요.").place(x = 105, y = 440)
        
        B_Frame = LabelFrame(self.window, text = '진행 상태',width = 355, height = 72)
        B_Frame['borderwidth'] = 2
        B_Frame['relief'] = 'ridge'
        B_Frame.place(x = 40, y = 450)
        
        B_Frame = Frame(self.window,width = 240, height = 30)
        B_Frame['borderwidth'] = 2
        B_Frame['relief'] = 'sunken'
        B_Frame.place(x = 65, y = 480)
        
        C_Frame = LabelFrame(self.window, text = 'ICSR/Safety 판단 결과',width = 355, height = 72)
        C_Frame['borderwidth'] = 2
        C_Frame['relief'] = 'ridge'
        C_Frame.place(x = 40, y = 540)
        
        C_Frame = Frame(self.window,width = 240, height = 30)
        C_Frame['borderwidth'] = 2
        C_Frame['relief'] = 'sunken'
        C_Frame.place(x = 65, y = 570)
        
        tk.Button(self.window, text = "시작", bg = '#448aff', fg = 'black',
                  font = ('sans 10 bold'), width = 9, command = self.search).place(x = 85, y = 640)
        tk.Button(self.window, text = "저장", bg = '#448aff', fg = 'black',
                  font = ('sans 10 bold'), width = 9, command = self.save_as).place(x = 273, y = 640)
        self.window.mainloop()
if __name__ == '__main__':
    
    if get_trained_source():
        search = LS_GUI(root)
        search.main_window()
    else:
        exit()
    print('done')
    #https://python-docx.readthedocs.io/en/latest/api/enum/WdColorIndex.html
    #https://ucwoogong.com/189 : progress bar
    #https://pythonguides.com/python-tkinter-progress-bar/ : progress bar2


# In[ ]:




